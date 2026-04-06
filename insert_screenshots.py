from docx import Document
from docx.shared import Inches
import os

doc_path = '/Users/kunaltailor/Desktop/Kunal/MPJ/VoteChain_Report_Final.docx'
screenshots_dir = '/Users/kunaltailor/Desktop/Kunal/MPJ/screenshots'

doc = Document(doc_path)

fig_map = {
    "Figure 1:": "fig01_architecture.png",
    "Figure 2:": "fig02_block_structure.png",
    "Figure 3:": "fig03_security_flow.png",
    "Figure 4:": "fig04_api_endpoints.png",
    "Figure 5:": "fig05_mongodb_schema.png",
    "Figure 6:": "fig06_home_screen.png",
    "Figure 7:": "fig07_voter_dashboard.png",
    "Figure 8:": "fig08_vote_casting.png",
    "Figure 9:": "fig09_vote_confirmation.png",
    "Figure 10:": "fig10_results_dashboard.png",
    "Figure 11:": "fig11_admin_dashboard.png",
    "Figure 12:": "fig06_home_voter_login.png", # Fallback for javafx
    "Figure 14:": "fig14_blockchain_validation.png",
    "Figure 15:": "fig15_results_admin.png"
}

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    for fig_key, img_name in fig_map.items():
        if text.startswith(fig_key):
            img_path = os.path.join(screenshots_dir, img_name)
            if os.path.exists(img_path):
                print(f"Adding {img_name} before {fig_key}")
                # We can insert image below the figure text or above it. 
                # Let's insert below.
                # Actually, adding above the caption is the standard format for figures.
                # Let's add above.
                p = para.insert_paragraph_before()
                p.alignment = 1 # Center align
                r = p.add_run()
                r.add_picture(img_path, width=Inches(6.0))
            else:
                print(f"Warning: {img_path} not found.")

doc.save('/Users/kunaltailor/Desktop/Kunal/MPJ/VoteChain_Report_Final.docx')
print("Saved VoteChain_Report_Final.docx")
